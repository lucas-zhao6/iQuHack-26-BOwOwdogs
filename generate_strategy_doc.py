#!/usr/bin/env python3
"""Generate the winning strategy document as a .docx file."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): color_hex,
        },
    )
    shading.append(shading_elem)


def add_code_block(doc, code_text):
    """Add a formatted code block paragraph."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.left_indent = Inches(0.3)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2E4057")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0F4F8")

    doc.add_paragraph()  # spacing


def build_document():
    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    # ============================================================
    # TITLE
    # ============================================================
    title = doc.add_heading("Comprehensive Winning Strategy: Circuit Fingerprint Challenge", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    subtitle = doc.add_paragraph("iQuHACK 2026 \u2014 Quantum Rings Track")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x5A, 0x7D, 0x9A)

    doc.add_paragraph()

    # ============================================================
    # TL;DR
    # ============================================================
    doc.add_heading("TL;DR", level=1)
    doc.add_paragraph(
        "Predict the minimum MPS threshold (discrete: 1..256) and forward runtime "
        "(continuous: seconds) for 18 hidden quantum circuits. Scoring is "
        "mean(threshold_score \u00d7 runtime_score) where under-predicting threshold = "
        "instant zero. 75% auto-grader, 25% presentation."
    )

    # ============================================================
    # SCORING ASYMMETRY
    # ============================================================
    doc.add_heading("Critical Scoring Asymmetry", level=1)
    doc.add_paragraph(
        "This asymmetry drives the entire strategy. The penalty structure is:"
    )

    add_table(
        doc,
        ["Scenario", "Score"],
        [
            ["Under-predict threshold by 1 step", "0.00 (catastrophic)"],
            ["Predict exactly right", "1.00"],
            ["Over-predict by 1 step", "0.50"],
            ["Over-predict by 2 steps", "0.25"],
            ["Runtime off by 2x either direction", "0.50"],
            ["Runtime off by 5x either direction", "0.20"],
        ],
    )

    p = doc.add_paragraph()
    run = p.add_run(
        "Key insight: It is ALWAYS better to over-predict threshold by 1 step (keep 50%) "
        "than risk under-predicting (lose 100%). We should systematically bias threshold "
        "predictions upward when uncertain."
    )
    run.bold = True

    # ============================================================
    # DATA PATTERNS
    # ============================================================
    doc.add_heading("What the Data Tells Us (patterns most teams will miss)", level=1)

    add_table(
        doc,
        ["Pattern", "Detail"],
        [
            [
                "68% of circuits need threshold \u2264 2",
                "But the 7% needing \u2265 32 are the ones that make or break scores",
            ],
            [
                "CPU beats GPU 85% of the time",
                "GPU setup overhead is 23x higher on average",
            ],
            [
                "Single precision is 2.5x faster",
                "With negligible fidelity loss for 93% of circuits",
            ],
            [
                "Circuit connectivity, not qubit count, predicts threshold",
                "GHZ at 130q needs threshold=2; QNN at 20q needs threshold=32",
            ],
            [
                "All-to-all connectivity = hard",
                "Portfolio, QNN, TwoLocalRandom, GroundState all have all-to-all CZ/CX and need high thresholds",
            ],
            [
                "2 precision anomalies exist",
                "GraphState CPU/double never converges; Shor CPU/single fails",
            ],
        ],
    )

    # ============================================================
    # THE 5 INNOVATIONS
    # ============================================================
    doc.add_heading("The 5 Innovations (What Sets Us Apart)", level=1)

    # --- Innovation 1 ---
    doc.add_heading("Innovation 1: MPS Cut-Pressure Profile (Physics-Informed Feature)", level=2)

    p = doc.add_paragraph()
    p.add_run("What everyone else will do: ").bold = True
    p.add_run("Count gates and qubits, maybe compute qubit degree.")

    p = doc.add_paragraph()
    p.add_run("What we do: ").bold = True
    p.add_run(
        "Compute the entanglement cut-pressure along the MPS bond chain. For each position k "
        "in the linear qubit ordering (positions 0..n-1), count how many two-qubit gates span "
        "across the cut at position k (i.e., gates connecting qubits i < k to qubits j \u2265 k). "
        "The maximum cut-pressure across all positions is a direct proxy for the maximum bond "
        "dimension the MPS needs \u2014 which is exactly what the threshold controls."
    )

    add_code_block(
        doc,
        'def compute_cut_pressure(two_qubit_gates, n_qubits):\n'
        '    """Physics-informed MPS difficulty estimator."""\n'
        '    pressure = [0] * n_qubits\n'
        '    for (q_i, q_j) in two_qubit_gates:\n'
        '        lo, hi = min(q_i, q_j), max(q_i, q_j)\n'
        '        for cut in range(lo + 1, hi + 1):\n'
        '            pressure[cut] += 1\n'
        '    return max(pressure), np.mean(pressure), np.std(pressure)',
    )

    p = doc.add_paragraph()
    p.add_run("Why it works: ").bold = True
    p.add_run(
        "MPS simulators factorize the state along a 1D chain. Long-range gates that cross many "
        "cuts create entanglement that requires larger bond dimensions. This feature directly "
        "captures what makes a circuit hard for the Quantum Rings simulator."
    )

    p = doc.add_paragraph()
    p.add_run("Demonstrable impact: ").bold = True
    p.add_run(
        "We can show that max_cut_pressure correlates more strongly with threshold than raw "
        "gate count or qubit count via an ablation study."
    )

    # --- Innovation 2 ---
    doc.add_heading(
        "Innovation 2: Asymmetric Threshold Prediction with Calibrated Safety Margin",
        level=2,
    )

    p = doc.add_paragraph()
    p.add_run("What everyone else will do: ").bold = True
    p.add_run("Standard multi-class classification, pick the argmax.")

    p = doc.add_paragraph()
    p.add_run("What we do: ").bold = True

    doc.add_paragraph(
        "1. Train an ordinal regression model on threshold rung index (0\u20138)",
        style="List Number",
    )
    doc.add_paragraph(
        "2. Compute prediction uncertainty via cross-validation residuals per family/config",
        style="List Number",
    )
    doc.add_paragraph(
        "3. Apply a calibrated upward bias: final_threshold = 2^(ceil(predicted_index + "
        "\u03b1 \u00d7 uncertainty)) where \u03b1 is tuned to maximize the actual scoring "
        "function (not accuracy)",
        style="List Number",
    )
    doc.add_paragraph(
        "4. Specifically optimize for the asymmetric scoring metric during hyperparameter search",
        style="List Number",
    )

    add_code_block(
        doc,
        "# Asymmetric loss that matches the actual competition scoring\n"
        "def competition_loss(y_true_idx, y_pred_idx):\n"
        "    if y_pred_idx < y_true_idx:\n"
        "        return 1.0  # catastrophic: score = 0\n"
        "    else:\n"
        "        steps_over = y_pred_idx - y_true_idx\n"
        "        return 1.0 - 2**(-steps_over)  # loss = 1 - score",
    )

    p = doc.add_paragraph()
    p.add_run("Demonstrable impact: ").bold = True
    p.add_run(
        "Show via leave-one-out CV that optimizing for the competition metric outperforms "
        "optimizing for classification accuracy."
    )

    # --- Innovation 3 ---
    doc.add_heading(
        "Innovation 3: Connectivity-Graph Fingerprinting (Circuit Family Detection)",
        level=2,
    )

    p = doc.add_paragraph()
    p.add_run("What everyone else will do: ").bold = True
    p.add_run(
        "Assume they know the family from the filename (but holdout filenames are mapped "
        "through an ID-map, so no family info!)."
    )

    p = doc.add_paragraph()
    p.add_run("What we do: ").bold = True
    p.add_run("Build a circuit family classifier based on structural fingerprints:")

    doc.add_paragraph("Gate type signature vector (presence/absence of rzz, cz, cp, ccx, etc.)", style="List Bullet")
    doc.add_paragraph("Connectivity pattern (nearest-neighbor vs star vs all-to-all)", style="List Bullet")
    doc.add_paragraph("Gate count ratios (two-qubit/total, parameterized/total)", style="List Bullet")
    doc.add_paragraph("Number of quantum registers (1 vs 2 vs 3)", style="List Bullet")

    doc.add_paragraph(
        "Then use the detected family as a powerful categorical feature. Each family has "
        "very consistent threshold behavior:"
    )

    add_table(
        doc,
        ["Family Group", "Expected Threshold", "Runtime Behavior"],
        [
            ["DJ, QFT, QPE, Grover", "1", "Scales with n_qubits\u00b2 (QFT) or gate_count"],
            ["GHZ, WState, CutBell", "2", "Very fast (linear scaling)"],
            ["VQE, QFTentangled", "2", "Moderate"],
            ["AE, GroundState, PricingCall, Shor", "4\u20138", "Moderate to high"],
            ["Portfolio*, QNN, TwoLocalRandom", "16\u2013256", "High (all-to-all kills MPS)"],
        ],
    )

    p = doc.add_paragraph()
    p.add_run("Demonstrable impact: ").bold = True
    p.add_run("Ablation showing that family-aware models outperform family-agnostic ones.")

    # --- Innovation 4 ---
    doc.add_heading("Innovation 4: Decomposed Runtime Model (Setup + Per-Shot)", level=2)

    p = doc.add_paragraph()
    p.add_run("What everyone else will do: ").bold = True
    p.add_run("Predict forward_wall_s as a single number.")

    p = doc.add_paragraph()
    p.add_run("What we do: ").bold = True
    p.add_run(
        "The data provides estimated_setup_s and estimated_per_shot_s. We model these separately:"
    )

    add_code_block(doc, "total_time = setup_time + per_shot_time * 10000")

    doc.add_paragraph("This decomposition is critical because:")

    doc.add_paragraph(
        "Setup time is dominated by backend (GPU setup is 23x slower) and circuit complexity",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Per-shot time is dominated by qubit count, threshold, and gate depth",
        style="List Bullet",
    )
    doc.add_paragraph(
        "They have completely different feature importance profiles",
        style="List Bullet",
    )

    doc.add_paragraph(
        "We train two regressors on log-scale, each with the most relevant features, and combine."
    )

    p = doc.add_paragraph()
    p.add_run("Demonstrable impact: ").bold = True
    p.add_run(
        "Show that decomposed prediction has lower error than direct prediction via "
        "RMSE comparison on CV."
    )

    # --- Innovation 5 ---
    doc.add_heading(
        "Innovation 5: Threshold Sweep Curve Shape Analysis (Training Enrichment)",
        level=2,
    )

    p = doc.add_paragraph()
    p.add_run("What everyone else will do: ").bold = True
    p.add_run("Only use the final selected threshold as the training label.")

    p = doc.add_paragraph()
    p.add_run("What we do: ").bold = True
    p.add_run(
        "The training data includes full threshold sweeps (fidelity at each rung 1..512). "
        "We extract the curve shape as additional training signal:"
    )

    doc.add_paragraph(
        "Convergence rate: How quickly does fidelity approach 1.0?",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Plateau detection: Does fidelity plateau below 0.99 (like GraphState/TwoLocalRandom)?",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Jump threshold: At which rung does fidelity make the biggest jump?",
        style="List Bullet",
    )

    doc.add_paragraph(
        "We use this to train a model that predicts the entire fidelity curve given circuit "
        "features, then derive the minimum threshold from the predicted curve. This provides "
        "richer supervision than just the final label."
    )

    # ============================================================
    # IMPLEMENTATION PLAN
    # ============================================================
    doc.add_heading("Implementation Plan", level=1)

    doc.add_heading("Phase 1: Data Pipeline & Feature Engineering", level=2)
    doc.add_paragraph("Parse all 36 QASM circuits \u2192 extract 30+ features", style="List Number")
    doc.add_paragraph("Load all 144 training results \u2192 build training dataframe", style="List Number")
    doc.add_paragraph("Implement MPS cut-pressure computation", style="List Number")
    doc.add_paragraph("Implement connectivity graph fingerprinting", style="List Number")
    doc.add_paragraph("Build circuit family classifier", style="List Number")

    doc.add_heading("Phase 2: Model Development", level=2)
    doc.add_paragraph(
        "Threshold prediction model (ordinal regression + asymmetric calibration)",
        style="List Number",
    )
    doc.add_paragraph(
        "Setup time prediction model (log-linear, backend-aware)", style="List Number"
    )
    doc.add_paragraph(
        "Per-shot time prediction model (log-linear, threshold-conditioned)",
        style="List Number",
    )
    doc.add_paragraph("Combine into final prediction pipeline", style="List Number")

    doc.add_heading("Phase 3: Validation & Tuning", level=2)
    doc.add_paragraph(
        "Leave-one-circuit-out cross-validation (critical: CV at circuit level, not result level)",
        style="List Number",
    )
    doc.add_paragraph(
        "Tune safety margin \u03b1 on the actual competition scoring function",
        style="List Number",
    )
    doc.add_paragraph("Ablation studies for each innovation", style="List Number")

    doc.add_heading("Phase 4: Submission", level=2)
    doc.add_paragraph("Build predict.py with CLI interface", style="List Number")
    doc.add_paragraph("Validate with provided scripts", style="List Number")
    doc.add_paragraph("Package submission", style="List Number")

    doc.add_heading("Phase 5: Presentation", level=2)
    doc.add_paragraph(
        "Ablation study showing each innovation\u2019s marginal contribution",
        style="List Number",
    )
    doc.add_paragraph(
        "Visualization of MPS cut-pressure vs threshold (compelling physics story)",
        style="List Number",
    )
    doc.add_paragraph(
        "Scoring metric analysis showing why asymmetric calibration matters",
        style="List Number",
    )

    # ============================================================
    # TECH STACK
    # ============================================================
    doc.add_heading("Tech Stack", level=1)

    doc.add_paragraph("numpy, pandas for data handling", style="List Bullet")
    doc.add_paragraph(
        "scikit-learn for models (GradientBoosting, RandomForest)", style="List Bullet"
    )
    doc.add_paragraph("networkx for connectivity graph analysis", style="List Bullet")
    doc.add_paragraph("re for QASM parsing", style="List Bullet")
    doc.add_paragraph("Standard Python for the predict.py pipeline", style="List Bullet")

    # ============================================================
    # RISK MITIGATION
    # ============================================================
    doc.add_heading("Risk Mitigation", level=1)

    add_table(
        doc,
        ["Risk", "Mitigation"],
        [
            [
                "Small training set (36 circuits, 144 results)",
                "Feature engineering over model complexity; physics-informed features reduce need for data",
            ],
            [
                "Unknown holdout circuit families",
                "Build robust family classifier; ensure features work well even without family labels",
            ],
            [
                "Threshold under-prediction",
                "Calibrated safety margin; never predict below what features suggest",
            ],
            [
                "Runtime spans 4 orders of magnitude",
                "Work in log-space; decompose into setup + per-shot",
            ],
            [
                "Anomalies (GraphState, Shor precision issues)",
                "Detect anomalies via feature outlier detection; handle gracefully",
            ],
        ],
    )

    # ============================================================
    # PRESENTATION STORY ARC
    # ============================================================
    doc.add_heading("Presentation Story Arc", level=1)

    story_points = [
        (
            "1. ",
            "\u201cWe identified that the scoring function is asymmetric \u2014 under-predicting "
            "threshold is catastrophic\u201d",
        ),
        (
            "2. ",
            "\u201cWe used physics-informed features: MPS cut-pressure directly predicts "
            "simulation difficulty\u201d",
        ),
        (
            "3. ",
            "\u201cWe decomposed runtime into setup + per-shot, modeling each with the right "
            "features\u201d",
        ),
        (
            "4. ",
            "\u201cWe built a circuit fingerprinting system that identifies algorithm families "
            "from structure alone\u201d",
        ),
        (
            "5. ",
            "\u201cOur calibrated safety margin optimizes for the actual competition metric, "
            "not surrogate accuracy\u201d",
        ),
    ]

    for num, text in story_points:
        p = doc.add_paragraph()
        p.add_run(num).bold = True
        p.add_run(text)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Ablation flow: ").bold = True
    p.add_run(
        "base model \u2192 +cut_pressure \u2192 +family_detection \u2192 "
        "+asymmetric_calibration \u2192 +decomposed_runtime \u2192 final score improvement"
    )

    # ============================================================
    # SAVE
    # ============================================================
    out_path = "/Users/langzhao/Documents/Projects/iQuHack-26/iQuHack-26-BOwOwdogs/Strategy_Circuit_Fingerprint_Challenge.docx"
    doc.save(out_path)
    print(f"Saved to {out_path}")


if __name__ == "__main__":
    build_document()
